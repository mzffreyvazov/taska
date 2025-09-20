# services/file_processor.py
"""File processing service for different document types"""
import os
import json
from typing import Optional, List, Dict
import pandas as pd
import pdfplumber
import PyPDF2
# PDF Libraries
PDF_LIBRARY = None
try:
    
    PDF_LIBRARY = 'pdfplumber'
except ImportError:
    try:
        PDF_LIBRARY = 'pypdf2'
    except ImportError:
        print("⚠️ PDF kitabxanası tapılmadı!")

import docx  # python-docx
import openpyxl

class FileProcessor:
    """Process different types of files and extract text"""
    
    def __init__(self):
        self.pdf_library = PDF_LIBRARY
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """Extract text from file based on type"""
        if not os.path.exists(file_path):
            return None
        
        extension = os.path.splitext(file_path)[1].lower()
        
        extractors = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.txt': self._extract_from_text,
            '.md': self._extract_from_text,
            '.json': self._extract_from_json,
            '.xlsx': self._extract_from_excel,
            '.xls': self._extract_from_excel
        }
        
        extractor = extractors.get(extension)
        if extractor:
            try:
                return extractor(file_path)
            except Exception as e:
                print(f"File extraction error ({extension}): {e}")
                return None
        
        return None
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        if not self.pdf_library:
            return "PDF kitabxanası yüklənməyib."
        
        if self.pdf_library == 'pdfplumber':
            return self._extract_with_pdfplumber(file_path)
        elif self.pdf_library == 'pypdf2':
            return self._extract_with_pypdf2(file_path)
        
        return ""
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """Extract with pdfplumber"""
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"\n=== Səhifə {page_num} ===\n{page_text}")
                
                # Extract tables
                tables = page.extract_tables()
                if tables:
                    text_parts.append(f"\n=== Səhifə {page_num} Cədvəlləri ===")
                    for table_idx, table in enumerate(tables, 1):
                        text_parts.append(f"\nCədvəl {table_idx}:")
                        for row in table:
                            if row and any(cell for cell in row if cell):
                                row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    def _extract_with_pypdf2(self, file_path: str) -> str:
        """Extract with PyPDF2"""
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"\n=== Səhifə {page_num + 1} ===\n{page_text}")
        
        return "\n".join(text_parts)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract from DOCX"""
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables, 1):
            text_parts.append(f"\n=== Cədvəl {table_idx} ===")
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    def _extract_from_text(self, file_path: str) -> str:
        """Extract from text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _extract_from_json(self, file_path: str) -> str:
        """Extract from JSON"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return self._json_to_text(data)
    
    def _json_to_text(self, obj, level=0) -> str:
        """Convert JSON object to text"""
        text_parts = []
        indent = "  " * level
        
        if isinstance(obj, dict):
            for key, value in obj.items():
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{indent}{key}:")
                    text_parts.append(self._json_to_text(value, level + 1))
                else:
                    text_parts.append(f"{indent}{key}: {value}")
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                if isinstance(item, (dict, list)):
                    text_parts.append(f"{indent}[{i}]:")
                    text_parts.append(self._json_to_text(item, level + 1))
                else:
                    text_parts.append(f"{indent}[{i}]: {item}")
        else:
            text_parts.append(f"{indent}{obj}")
        
        return "\n".join(text_parts)
    
    def _extract_from_excel(self, file_path: str) -> str:
        """Extract from Excel files"""
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            text_parts = []
            
            for sheet_name, df in excel_data.items():
                if df.empty:
                    continue
                
                text_parts.append(f"\n=== {sheet_name} Vərəqi ===")
                
                # Clean NaN values
                df = df.fillna('')
                
                # Headers
                headers = " | ".join([str(col) for col in df.columns])
                text_parts.append(f"Başlıqlar: {headers}")
                text_parts.append("-" * min(80, len(headers)))
                
                # Data (first 100 rows)
                max_rows = min(100, len(df))
                for idx, row in df.head(max_rows).iterrows():
                    row_text = " | ".join([str(val) if val != '' else "boş" for val in row])
                    text_parts.append(f"Sətir {idx + 1}: {row_text}")
                
                if len(df) > 100:
                    text_parts.append(f"... və daha {len(df) - 100} sətir")
                
                # Statistics for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if not numeric_cols.empty:
                    text_parts.append("\nRəqəmsal Statistika:")
                    for col in numeric_cols:
                        col_data = df[col].dropna()
                        if not col_data.empty:
                            text_parts.append(
                                f"{col}: Min={col_data.min()}, Max={col_data.max()}, "
                                f"Orta={col_data.mean():.2f}"
                            )
            
            return "\n".join(text_parts)
            
        except Exception as e:
            print(f"Excel extraction error: {e}")
            return ""
    
    def get_file_type(self, filename: str) -> str:
        """Get file type from filename"""
        ext = os.path.splitext(filename)[1].lower()
        type_map = {
            '.pdf': 'PDF',
            '.docx': 'Word',
            '.txt': 'Text',
            '.md': 'Markdown',
            '.json': 'JSON',
            '.xlsx': 'Excel',
            '.xls': 'Excel'
        }
        return type_map.get(ext, 'Unknown')
    
    def validate_file(self, filename: str, max_size: int) -> tuple[bool, str]:
        """Validate file type and size"""
        from config import Config
        
        ext = os.path.splitext(filename)[1].lower()
        
        if ext not in Config.SUPPORTED_EXTENSIONS:
            return False, f"Dəstəklənməyən fayl tipi. Dəstəklənənlər: {', '.join(Config.SUPPORTED_EXTENSIONS)}"
        
        return True, ""